from PyQt6 import QtWidgets, uic
from PyQt6.uic import loadUi
from PyQt6.QtWidgets import *
from PyQt6.QtWidgets import QVBoxLayout, QFileDialog, QApplication, QWidget, QPushButton, QProgressBar, QMessageBox, QMainWindow, QApplication
from PyQt6.QtGui import QAction
from PyQt6.QtCore import QSize, Qt, QThread, pyqtSignal
from pathlib import Path
import posixpath
import edge_tts

import os
import asyncio
from docx import Document
import sys
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import random
import threading

import asyncio
import random
import shutil
import wave
import audioread

import re

class AudioWorker(QThread):
    progress_signal = pyqtSignal(int)
    completed_signal = pyqtSignal()

    def __init__(self, word_pairs, path, Eng_repeat, Ch_repeat, silence_between_repeats, silence_between_innter_repeats):
        super().__init__()
        self.word_pairs = word_pairs
        self.path = path
        self.Eng_repeat = Eng_repeat
        self.Ch_repeat = Ch_repeat
        self.silence_between_repeats = silence_between_repeats
        self.silence_between_innter_repeats = silence_between_innter_repeats

    def run(self):
        self.audio(self.word_pairs, self.path, self.progress_signal.emit, self.Eng_repeat, self.Ch_repeat, self.silence_between_repeats, self.silence_between_innter_repeats)
        self.completed_signal.emit()

    def audio(self, word_pairs, path, progress_callback, Eng_repeat, Ch_repeat, silence_between_repeats, silence_between_innter_repeats):
        async def create_tts_audio(text, name, voice, max_retries=50):
            #ShortText = ((((text.replace(" ", "_")).replace(",", "_")).replace("'", "_")).replace('"', '_')).replace(".", "_")
            #max_filename_length = 100
            #truncated_text = ShortText[:max_filename_length]
            filename = f"{name}.mp3"
            retries = 0
            while retries < max_retries:
                try:
                    # Generate TTS audio using edge_tts (same as before)
                    communicate = edge_tts.Communicate(text, voice=voice)
                    await communicate.save(filename)
                    return filename
                except Exception as e:
                    print(f"Error generating TTS: {e}")
                    await asyncio.sleep(2 + random.uniform(0, 2))
                    retries += 1
            raise Exception(f"Failed to generate TTS for '{text}' after {max_retries} attempts.")

        async def generate_audio(word_pairs, path, progress_callback, Eng_repeat, Ch_repeat, silence_between_repeats, silence_between_innter_repeats):
            final_audio_files = []  # Store the MP3 files for combining later
            total_word_pairs = len(word_pairs)

            for index, (eng, chi) in enumerate(word_pairs):
                eng_filename = await create_tts_audio(eng, str(index), voice='en-US-JennyNeural')
                chi_filename = await create_tts_audio(chi, str(str(index)+"second"), voice='zh-CN-YunyangNeural')

                # Repeat English and Chinese
                for _ in range(Eng_repeat):
                    final_audio_files.append(eng_filename)
                    final_audio_files.append(None)  # Represent silence
                for _ in range(Ch_repeat):
                    final_audio_files.append(chi_filename)
                    final_audio_files.append(None)  # Represent silence

                # Add silence between repeats
                final_audio_files.append(None)

                # Update progress
                progress = int((index + 1) * 100 / total_word_pairs)
                progress_callback(progress)


            final_audio = self.merge_audio_files(final_audio_files)


            # Save the final audio as an MP3
            with open(path, 'wb') as out_file:
                out_file.write(final_audio)

            for i in range(0, len(final_audio_files)):
                if str(final_audio_files[i]) != "None":
                    os.remove(str(final_audio_files[i]))
                else:
                    pass

        # Generate the final audio

        final_audio = asyncio.run(generate_audio(word_pairs, path, progress_callback, Eng_repeat, Ch_repeat, silence_between_repeats, silence_between_innter_repeats))

    def merge_audio_files(self, audio_files):
        final_audio = b''  # Start with an empty byte string

        for file in audio_files:
            if file:
                # Read the MP3 file and append its content
                with open(file, 'rb') as f:
                    audio_content = f.read()
                    final_audio += audio_content
            else:
                # Add silence (if None) -- simply add empty byte for now
                silence_content = b'\0' * (self.silence_between_innter_repeats * 1000)  # Adjust silence duration
                final_audio += silence_content

        return final_audio


#init################################################################
global source_type, path1, output_path, option, muti_word_list
source_type = ""
path1 = ""
output_path = ""
option = 0
muti_word_list = []
def find_files(path, extension):
    files_list = []
    for file in os.listdir(path):
        if file.endswith(extension):
            files_list.append(Path(os.path.join(path, file)))
    return files_list
if hasattr(sys, '_MEIPASS'):
    # When running the PyInstaller-extracted binary
    base_path = sys._MEIPASS
else:
    # When running the script normally
    base_path = os.path.dirname(os.path.abspath(__file__))
######################################################################

class MainWindow(QMainWindow):
    def __init__(self):
        super(MainWindow, self).__init__()
        loadUi(os.path.join(base_path, '1.ui'), self)

        self.current_value = 0
        self.actionpdf.triggered.connect(self.page2)
        self.pushbutton_start.clicked.connect(self.start)
        self.pushbutton_source.clicked.connect(self.select_input_path)
        self.pushbutton_choose.clicked.connect(self.select_output_path)
        self.spinBox_7.setValue(200)

    def start(self):
        global double, option
        if path1 == "":
            QMessageBox.critical(self, '错误', '未选择源文件/文件夹，请重新选择 E:0001')
            return

        if output_path == "":
            QMessageBox.critical(self, '错误', '未选择目标文件/文件夹，请重新选择 E:0002')
            return

        if self.radioButton_all.isChecked():
            files_list = find_files(path1, ".docx")
            option = 2

            if self.checkBox.isChecked() and self.checkBox_2.isChecked():
                double = True
                en1 = self.spinBox.value()
                cn1 = self.spinBox_2.value()
                en2 = self.spinBox_3.value()
                cn2 = self.spinBox_4.value()
            elif self.checkBox.isChecked() and not self.checkBox_2.isChecked():
                double = False
                en1 = self.spinBox.value()
                cn1 = self.spinBox_2.value()
                en2 = None
                cn2 = None
            else:
                QMessageBox.critical(self, '错误', 'docx文件需要指定单词列位置 E:0003')
                return

            # Process multiple files
            for file in files_list:
                word_list = self.word_list("docx", posixpath.join(path1, Path(file.name)), double, en1, cn1, en2, cn2)
                output_file = Path(output_path, f"{file.stem}.mp3")

                # Create a worker thread to generate the audio
                self.audio_thread = AudioWorker(word_list, output_file, self.spinBox_5.value(), self.spinBox_6.value(), self.spinBox_7.value(), self.spinBox_8.value())
                self.audio_thread.progress_signal.connect(self.update_progress_bar)
                self.audio_thread.completed_signal.connect(self.audio_completed)

                self.audio_thread.start()

        elif self.radioButton_single.isChecked():
            option = 1
            if self.checkBox.isChecked() and self.checkBox_2.isChecked():
                double = True
                en1 = self.spinBox.value()
                cn1 = self.spinBox_2.value()
                en2 = self.spinBox_3.value()
                cn2 = self.spinBox_4.value()
            elif self.checkBox.isChecked() and not self.checkBox_2.isChecked():
                double = False
                en1 = self.spinBox.value()
                cn1 = self.spinBox_2.value()
                en2 = None
                cn2 = None
            else:
                QMessageBox.critical(self, '错误', 'docx文件需要指定单词列位置 E:0005')
                return

            # Single file processing
            word_pairs = self.word_list(source_type, path1, double, en1, cn1, en2, cn2)

            # Start the worker thread for single file
            self.audio_thread = AudioWorker(word_pairs, output_path, self.spinBox_5.value(), self.spinBox_6.value(), self.spinBox_7.value(), self.spinBox_8.value())
            self.audio_thread.progress_signal.connect(self.update_progress_bar)
            self.audio_thread.completed_signal.connect(self.audio_completed)

            self.audio_thread.start()

    def create_docx(self, text, file_path):
        doc = Document()
        
        # Split text into paragraphs if there are multiple lines
        paragraphs = text.split('\n')
        
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)

        doc.save(file_path)

    def audio_completed(self):
        QMessageBox.information(self, '提示', '所有音频已生成')

    def word_list(self, type, path, double, en1, cn1, en2, cn2):
            word_pairs = []
            if type == "docx":

                word_pairs = self.extract_word_pairs(path, double, en1, cn1, en2, cn2)
            elif type == "txt":

                with open(path, "r", encoding="utf-8") as f:
                    lines = [line.strip() for line in f.readlines()]
                word_pairs = [(lines[i], lines[i + 1]) for i in range(0, len(lines), 2)]
            return word_pairs

    def select_output_path(self):
        if self.radioButton_single.isChecked() == True:
            self.select_output_file()

        elif self.radioButton_all.isChecked() == True:
            self.select_output_folder()

        elif self.radioButton_single.isChecked() == False and self.radioButton_all.isChecked() == False:
            QMessageBox.critical(
                self,
                '错误',
                '请选择操作数量 E:0008'
                )
        else:
            QMessageBox.critical(
                self,
                '未知错误',
                '请联系开发者 E:0009'
                )

    def extract_word_pairs(self, doc_path, double, en1, cn1, en2, cn2):
        doc = Document(doc_path)
        word_pairs = []

        for table in doc.tables:
            for row in table.rows:
                left_english = row.cells[en1-1].text.strip()
                left_chinese = row.cells[cn1-1].text.strip()
                if left_english and left_chinese:
                    word_pairs.append((left_english, left_chinese))
            if double == True:
                for row in table.rows:
                    right_english = row.cells[en2-1].text.strip()
                    right_chinese = row.cells[cn2-1].text.strip()
                    if right_english and right_chinese:
                        word_pairs.append((right_english, right_chinese))

        return word_pairs
                
    def select_output_file(self):
        global output_path
        dirname, _ = QFileDialog.getSaveFileName(
            self,
            "Save File",
            "",
            "All Files (*)",
        )

        if not dirname.endswith(".mp3"):
            dirname += ".mp3"
        output_path = Path(dirname)
        return

    def select_output_folder(self):
        global output_path
        output_path = Path(QFileDialog.getExistingDirectory(self, "Select Directory"))

        return

    def select_input_path(self):
        if self.radioButton_single.isChecked() == True:
            self.select_file()

        elif self.radioButton_all.isChecked() == True:
            self.select_in_folder()

        elif self.radioButton_single.isChecked() == False and self.radioButton_all.isChecked() == False:
            QMessageBox.critical(
                self,
                '错误',
                '请选择操作数量 E:0008'
                )
        else:
            QMessageBox.critical(
                self,
                '未知错误',
                '请联系开发者 E:0009'
                )
        
    def select_in_folder(self):
        global path1
        path1 = Path(QFileDialog.getExistingDirectory(self, "Select Directory"))
        if path1:
            pass
        else:
            QMessageBox.critical(
                self,
                '错误',
                '未选中 E:0010'
                )
            return

    def select_file(self):
        global path1, source_type
        filename, _ = QFileDialog.getOpenFileName(self, "Select a File")
        if filename:
            if filename.endswith(".txt"):
                source_type = "txt"
                path1 = Path(filename)

            elif filename.endswith(".docx"):
                source_type = "docx"

                path1 = Path(filename)
            else:
                QMessageBox.critical(
                    self,
                    '错误',
                    '选中的文件非docx或者txt，请重新选择 E:0011'
                    )
                return
        else:
            QMessageBox.critical(
                self,
                '错误',
                '未选中 E:0012'
                )       
            return
           

    def page2(self):
        widget.setCurrentIndex(1)   

    def update_progress_bar(self, value):
        # Set the progress bar directly without accumulation
        self.progressBar.setValue(value)

#init########################
output_path2 = ""
path2 = ""
source_type2 = ""
#############################

class PDFWorker(QThread):
    progress_signal = pyqtSignal(int)
    completed_signal = pyqtSignal()

    def __init__(self, files_list, output_path2, start_page, end_page, spinBox_2_value):
        super().__init__()
        self.files_list = files_list
        self.output_path2 = output_path2
        self.start_page = start_page
        self.end_page = end_page
        self.spinBox_2_value = spinBox_2_value

    def run(self):
        # Calculate the total number of pages across all PDFs for progress tracking
        total_pages = 0
        for file in self.files_list:
            doc = fitz.open(file)
            total_pages += (len(doc) - self.spinBox_2_value) - self.start_page
            doc.close()

        pages_processed = 0

        def process_file(file, index):
            nonlocal pages_processed  # Allow updating `pages_processed` in the nested function
            doc = fitz.open(file)
            num_pages_in_doc = len(doc)  # Get total pages for this document

            # Calculate the end_page for the current document, ensuring it doesn't exceed the total pages
            end_page = min(num_pages_in_doc - self.spinBox_2_value, self.end_page)

            text = ""

            
            for page_num in range(self.start_page, end_page):
                page = doc.load_page(page_num)
                pix = page.get_pixmap()
                image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                text += "\n" + pytesseract.image_to_string(image)
                pages_processed += 1

                # Emit progress signal
                progress = int((pages_processed / total_pages) * 100)
                self.progress_signal.emit(progress)

            doc.close()

            # Ensure each output file has a unique name
            output_file = Path(self.output_path2, f"{Path(file).stem}_{index+1}.docx")
            self.create_docx(text, output_file)

        # Process each file using a separate thread
        threads = []
        for i, file in enumerate(self.files_list):
            thread = threading.Thread(target=process_file, args=(file, i))
            threads.append(thread)
            thread.start()

        for thread in threads:
            thread.join()

        # Emit completed signal
        self.completed_signal.emit()




    def create_docx(self, text, file_path):
        doc = Document()
        paragraphs = text.split('\n')
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)
        doc.save(file_path)


class window2(QMainWindow):
    def __init__(self):
        super(window2, self).__init__()
        loadUi(os.path.join(base_path, '2.ui'), self)

        self.action.triggered.connect(self.page1)
        self.pushButton_start.clicked.connect(self.start)
        self.pushButton.clicked.connect(self.select_input_path)
        self.pushButton_2.clicked.connect(self.select_output_folder)

    def start(self):
        global path2, output_path2

        if path2 == "":
            QMessageBox.critical(self, '错误', '未选择源文件/文件夹，请重新选择 E:0014')
            return

        if output_path2 == "":
            QMessageBox.critical(self, '错误', '未选择目标文件/文件夹，请重新选择 E:0015')
            return

        start_page = self.spinBox.value()
        spinBox_2_value = self.spinBox_2.value()

        if self.radioButton_single.isChecked():
            files_list = [path2]  # Single file
        elif self.radioButton_all.isChecked():
            files_list = find_files(path2, ".pdf")  # Multiple files

        # Open the first file to calculate the total number of pages
        doc = fitz.open(files_list[0])
        end_page = len(doc)  # Set end_page to the total number of pages in the PDF
        doc.close()

        # Create a PDF worker thread and start processing
        self.pdf_worker = PDFWorker(files_list, output_path2, start_page, end_page, spinBox_2_value)
        self.pdf_worker.progress_signal.connect(self.update_progress_bar)
        self.pdf_worker.completed_signal.connect(self.pdf_processing_completed)
        self.pdf_worker.start()


    def update_progress_bar(self, value):
        self.progressBar.setValue(value)

    def pdf_processing_completed(self):
        QMessageBox.information(self, '提示', '所有pdf已生成')



    def select_output_path(self):
        if self.radioButton_single.isChecked() == True:
            self.select_output_file()

        elif self.radioButton_all.isChecked() == True:
            self.select_output_folder()

        elif self.radioButton_single.isChecked() == False and self.radioButton_all.isChecked() == False:
            QMessageBox.critical(
                self,
                '错误',
                '请选择操作数量 E:0008'
                )
        else:
            QMessageBox.critical(
                self,
                '未知错误',
                '请联系开发者 E:0009'
                )
                
    def select_output_file(self):
        global output_path2
        dirname, _ = QFileDialog.getSaveFileName(
            self,
            "Save File",
            "",
            "All Files (*)",
        )
        if not dirname.endswith(".docx"):
            dirname += ".docx"
        output_path2 = Path(dirname)

    def select_output_folder(self):
        global output_path2
        output_path2 = Path(QFileDialog.getExistingDirectory(self, "Select Directory"))

        return

    def select_input_path(self):
        if self.radioButton_single.isChecked():
            self.select_file()
        elif self.radioButton_all.isChecked():
            self.select_in_folder()
        elif not self.radioButton_single.isChecked() and not self.radioButton_all.isChecked():
            QMessageBox.critical(
                self,
                '错误',
                '请选择操作数量 E:0016'
            )
        else:
            QMessageBox.critical(
                self,
                '未知错误',
                '请联系开发者 E:0017'
            )

    def select_in_folder(self):
        global path2
        path2 = Path(QFileDialog.getExistingDirectory(self, "Select Directory"))
        if not path2:
            QMessageBox.critical(
                self,
                '错误',
                '未选中 E:0018'
            )

    def select_file(self):
        global path2, source_type2
        filename, _ = QFileDialog.getOpenFileName(self, "Select a File")
        if filename:
            if filename.endswith(".pdf"):
                source_type2 = "pdf"
            else:
                QMessageBox.critical(
                    self,
                    '错误',
                    '选中的文件非docx或者txt，请新选择 E:0019'
                )
                return
            path2 = Path(filename)

    def page1(self):
        widget.setCurrentIndex(0)

 


if __name__ == '__main__':
    app = QApplication(sys.argv)
    app.setStyle("fusion")
    widget = QtWidgets.QStackedWidget()
    mainwindow = MainWindow()
    window2 = window2()
    widget.addWidget(mainwindow)
    widget.addWidget(window2)
    widget.setFixedHeight(640)
    widget.setFixedWidth(480)
    widget.setWindowTitle("生成器 V2.0")
    widget.show()
    sys.exit(app.exec())
